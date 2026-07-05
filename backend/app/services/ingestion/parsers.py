"""File parsers — extract raw text + metadata from uploaded documents.

Each parser accepts the file's bytes and returns a :class:`ParsedDocument`.
Parsers only *extract*; cleaning/normalisation happens in
:mod:`app.services.ingestion.cleaning` and chunking in
:mod:`app.services.ingestion.chunking`.
"""

from __future__ import annotations

import csv
import io
from dataclasses import dataclass, field
from pathlib import PurePosixPath, PureWindowsPath
from typing import Any, Protocol

from app.services.ingestion.ocr import get_ocr_engine


class ParserError(Exception):
    """Raised when a file cannot be parsed."""


class UnsupportedTypeError(ParserError):
    """Raised when no parser exists for a file type."""


@dataclass
class ParsedDocument:
    """The result of parsing a file."""

    text: str
    metadata: dict[str, Any] = field(default_factory=dict)
    page_count: int | None = None
    # Raw per-page texts (when the format has pages, e.g. PDF) so chunks can
    # carry page numbers into the vector index.
    pages: list[str] | None = None


class Parser(Protocol):
    """Interface every parser implements."""

    doc_types: tuple[str, ...]

    def parse(self, content: bytes, filename: str) -> ParsedDocument: ...


# --------------------------------------------------------------------------- #
# Plain text / markdown
# --------------------------------------------------------------------------- #
class TxtParser:
    """Plain-text and markdown files."""

    doc_types = ("txt", "md")

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        text = _decode_bytes(content)
        return ParsedDocument(
            text=text,
            metadata={"encoding_note": "decoded as utf-8 (with fallbacks)"},
        )


# --------------------------------------------------------------------------- #
# CSV
# --------------------------------------------------------------------------- #
class CsvParser:
    """CSV files — rendered row-by-row with column context preserved."""

    doc_types = ("csv",)

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        raw = _decode_bytes(content)
        sample = raw[:8192]
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel  # default comma dialect

        reader = csv.reader(io.StringIO(raw), dialect)
        rows = [row for row in reader if any(cell.strip() for cell in row)]
        if not rows:
            raise ParserError("CSV file contains no data")

        header = [cell.strip() for cell in rows[0]]
        data_rows = rows[1:]

        # Render each record as "Column: value" pairs so every chunk keeps
        # its column context after splitting.
        lines: list[str] = []
        for row in data_rows:
            pairs = [
                f"{header[i] if i < len(header) else f'column_{i + 1}'}: {cell.strip()}"
                for i, cell in enumerate(row)
                if cell.strip()
            ]
            if pairs:
                lines.append(" | ".join(pairs))

        text = "\n".join(lines) if lines else " | ".join(header)
        return ParsedDocument(
            text=text,
            metadata={
                "columns": header,
                "column_count": len(header),
                "row_count": len(data_rows),
                "delimiter": dialect.delimiter,
            },
        )


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
class PdfParser:
    """PDF files via :mod:`pypdf` (text-based PDFs; scanned PDFs need OCR)."""

    doc_types = ("pdf",)

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover
            raise ParserError("pypdf is not installed") from exc

        try:
            reader = PdfReader(io.BytesIO(content))
            if reader.is_encrypted:
                # Try the empty password; many "encrypted" PDFs open with it.
                try:
                    reader.decrypt("")
                except Exception as exc:
                    raise ParserError("PDF is password-protected") from exc

            pages = [page.extract_text() or "" for page in reader.pages]
        except ParserError:
            raise
        except Exception as exc:
            raise ParserError(f"Could not read PDF: {exc}") from exc

        text = "\n\n".join(page for page in pages if page.strip())
        if not text.strip():
            raise ParserError(
                "No extractable text found in PDF — it may be a scanned "
                "document that requires OCR."
            )

        metadata: dict[str, Any] = {}
        try:
            info = reader.metadata
            if info:
                for key, meta_key in (
                    ("title", "/Title"),
                    ("author", "/Author"),
                    ("subject", "/Subject"),
                    ("creator", "/Creator"),
                    ("producer", "/Producer"),
                ):
                    value = info.get(meta_key)
                    if value:
                        metadata[key] = str(value)
        except Exception:  # metadata is best-effort
            pass

        return ParsedDocument(
            text=text, metadata=metadata, page_count=len(reader.pages), pages=pages
        )


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
class DocxParser:
    """Word documents via :mod:`python-docx` — paragraphs and tables."""

    doc_types = ("docx",)

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        try:
            import docx
        except ImportError as exc:  # pragma: no cover
            raise ParserError("python-docx is not installed") from exc

        try:
            document = docx.Document(io.BytesIO(content))
        except Exception as exc:
            raise ParserError(f"Could not read DOCX: {exc}") from exc

        parts: list[str] = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n".join(parts)
        if not text.strip():
            raise ParserError("DOCX file contains no readable text")

        metadata: dict[str, Any] = {}
        try:
            props = document.core_properties
            if props.title:
                metadata["title"] = props.title
            if props.author:
                metadata["author"] = props.author
            if props.subject:
                metadata["subject"] = props.subject
            if props.created:
                metadata["created"] = props.created.isoformat()
            if props.last_modified_by:
                metadata["last_modified_by"] = props.last_modified_by
        except Exception:
            pass

        return ParsedDocument(text=text, metadata=metadata)


# --------------------------------------------------------------------------- #
# Images (OCR)
# --------------------------------------------------------------------------- #
class ImageParser:
    """Images via the pluggable OCR architecture (see :mod:`.ocr`)."""

    doc_types = ("image",)

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        engine = get_ocr_engine()
        if not engine.is_available():
            raise ParserError(
                "Image uploaded but no OCR engine is installed. Install the "
                "optional OCR extras (pytesseract + Pillow + Tesseract) and "
                "re-index this document."
            )
        try:
            text = engine.extract_text(content)
        except Exception as exc:
            raise ParserError(f"OCR failed: {exc}") from exc
        if not text.strip():
            raise ParserError("OCR found no text in the image")
        return ParsedDocument(text=text, metadata={"ocr_engine": engine.name})


# --------------------------------------------------------------------------- #
# Registry
# --------------------------------------------------------------------------- #
_PARSERS: list[Parser] = [TxtParser(), CsvParser(), PdfParser(), DocxParser(), ImageParser()]

SUPPORTED_DOC_TYPES: tuple[str, ...] = tuple(
    doc_type for parser in _PARSERS for doc_type in parser.doc_types
)

_EXTENSION_MAP: dict[str, str] = {
    "pdf": "pdf",
    "docx": "docx",
    "txt": "txt",
    "text": "txt",
    "log": "txt",
    "md": "md",
    "markdown": "md",
    "csv": "csv",
    "png": "image",
    "jpg": "image",
    "jpeg": "image",
    "webp": "image",
    "bmp": "image",
    "tif": "image",
    "tiff": "image",
}

_MIME_MAP: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/markdown": "md",
    "text/csv": "csv",
    "application/csv": "csv",
}


def detect_doc_type(filename: str, mime_type: str | None = None) -> str:
    """Infer the document type from the file extension (mime as fallback)."""
    # Handle both windows and posix style client paths
    name = PureWindowsPath(PurePosixPath(filename).name).name
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    if ext in _EXTENSION_MAP:
        return _EXTENSION_MAP[ext]
    if mime_type:
        base_mime = mime_type.split(";")[0].strip().lower()
        if base_mime in _MIME_MAP:
            return _MIME_MAP[base_mime]
        if base_mime.startswith("image/"):
            return "image"
        if base_mime.startswith("text/"):
            return "txt"
    raise UnsupportedTypeError(
        f"Unsupported file type '{ext or mime_type or 'unknown'}'. "
        f"Supported: pdf, docx, txt, md, csv, images (png/jpg/webp/tiff)."
    )


def get_parser(doc_type: str) -> Parser:
    """Return the parser responsible for the given document type."""
    for parser in _PARSERS:
        if doc_type in parser.doc_types:
            return parser
    raise UnsupportedTypeError(f"No parser registered for type '{doc_type}'")


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #
def _decode_bytes(content: bytes) -> str:
    """Decode bytes to text, trying UTF-8 variants before latin-1."""
    for encoding in ("utf-8-sig", "utf-8"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("latin-1", errors="replace")
